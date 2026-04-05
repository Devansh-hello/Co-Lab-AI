#!/usr/bin/env python3
"""Generate Co-Lab AI Major Project Report (.docx) with proper formatting."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Global defaults ───────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Heading 1 style
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(12)
h1.paragraph_format.line_spacing = 1.5

# Heading 2 style
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(8)
h2.paragraph_format.line_spacing = 1.5

# Heading 3 style
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(6)
h3.paragraph_format.line_spacing = 1.5

# ─── Margins (1 inch) ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ─── Helper functions ──────────────────────────────────────────────

def add_page_number(doc):
    """Add page numbers to footer (bottom center)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        run3._r.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)


def para(text, bold=False, italic=False, align=None, size=None, indent=False):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size or 12)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p


def numbered(text, level=0):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_table(headers, rows, caption=None, table_num=None):
    """Add a formatted table with optional caption."""
    if caption and table_num:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Table {table_num}: {caption}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()  # spacing after table
    return table


def add_figure_caption(fig_num, caption):
    """Add figure caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Figure {fig_num}: {caption}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    return p


def section_break():
    """Add a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
para('CO-LAB AI', bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('The Multi-Agent AI Platform That Builds', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Full-Stack Web Applications', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
para('A Major Project Report', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Submitted in partial fulfilment of the requirements for the degree of', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Bachelor of Technology', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para('in', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Computer Science and Engineering', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
para('Submitted by:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Devansh', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Palak', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
para('Department of Computer Science and Engineering', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('2025', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  PROJECT COMPLETION CERTIFICATE (page ii)
# ═══════════════════════════════════════════════════════════════════

doc.add_paragraph()
para('PROJECT COMPLETION CERTIFICATE', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

para('This is to certify that the project entitled "Co-Lab AI: The Multi-Agent AI Platform That Builds Full-Stack Web Applications" is a bonafide work carried out by Devansh and Palak in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.', indent=True)
doc.add_paragraph()
para('The project work has been carried out under my supervision and guidance. The work embodied in this project has not been submitted for the award of any other degree.', indent=True)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

para('Date: _______________', size=12)
doc.add_paragraph()
para('Project Guide:', size=12)
para('Name: ___________________________', size=12)
para('Designation: _____________________', size=12)
para('Signature: _______________________', size=12)
doc.add_paragraph()
doc.add_paragraph()
para('Head of Department:', size=12)
para('Name: ___________________________', size=12)
para('Signature: _______________________', size=12)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  INDEX PAGE
# ═══════════════════════════════════════════════════════════════════

para('INDEX', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ['S. No.', 'Content', 'Page No.'],
    [
        ['1.', 'Project Completion Certificate', 'ii'],
        ['2.', 'Abstract', '1'],
        ['3.', 'Introduction', '2-5'],
        ['4.', 'Motivation', '6-7'],
        ['5.', 'Literature Review / Comparative Work Evaluation', '8-10'],
        ['6.', 'Gap Analysis', '11-13'],
        ['7.', 'Problem Statement', '14'],
        ['8.', 'Objectives', '15-16'],
        ['9.', 'Methodology', '17-24'],
        ['10.', 'Tools / Platform Used', '25-26'],
        ['11.', 'Environmental Setup', '27-28'],
        ['12.', 'Implementation', '29-32'],
        ['13.', 'Results and Discussion', '33-35'],
        ['14.', 'Conclusion & Future Work', '36-37'],
        ['15.', 'References', '38'],
        ['16.', 'Annexure I: Plagiarism Declaration Certificate', '39'],
        ['17.', 'Annexure II: Plagiarism Report', '40'],
    ]
)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  1. ABSTRACT (page 1)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Abstract', level=1)

para('Co-Lab AI is a research-driven, multi-agent AI coding platform that generates complete, working full-stack web applications from natural language descriptions. Unlike conventional single-model AI coding assistants that produce unvalidated code fragments, Co-Lab AI employs a coordinated pipeline of six specialized AI agents, each responsible for a distinct phase of the software development lifecycle: requirement understanding, architectural planning, frontend development, backend development, code review, and independent test generation. A supervisory quality scoring function and an automated feedback loop ensure that the generated output meets measurable quality standards before delivery to the user.', indent=True)

doc.add_paragraph()

para('The platform is built upon techniques drawn from 19 peer-reviewed research papers published at top-tier AI and software engineering venues including ICLR, NeurIPS, ACL, EMNLP, and FSE. Key innovations include a formal API contract specification that guarantees frontend-backend compatibility, independent test generation that prevents implementation bias (inspired by AgentCoder), trajectory reduction for token efficiency (inspired by AgentDiet), chain-of-verification for hallucination reduction, and selective agent re-execution during the feedback loop (inspired by AgentDropout).', indent=True)

doc.add_paragraph()

para('The system is implemented as a monorepo with a React 19 / Next.js 16 frontend featuring an in-browser IDE with live preview via WebContainers, and a Node.js / Express 5 backend with real-time WebSocket streaming. The AI infrastructure supports five providers (OpenAI, Anthropic, Google Gemini, OpenRouter, and GLM), allowing users to select different models for different agent roles. Evaluation across three benchmark scenarios demonstrated quality grades of A (90/100) for simple applications, C (78/100) for complex multi-feature applications, and B (87/100) for iterative feature additions, with token efficiency improvements of 35-45% over the unoptimized baseline.', indent=True)

doc.add_paragraph()

para('Keywords: Multi-Agent Systems, Code Generation, Large Language Models, Software Engineering, API Contract Specification, Quality Assurance, WebContainers', bold=True, italic=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  2. INTRODUCTION (pages 2-5)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Background', level=2)

para('The rapid advancement of Large Language Models (LLMs) has transformed the landscape of software development. Tools such as GitHub Copilot, ChatGPT, and Claude have demonstrated the ability to generate code from natural language descriptions, leading to a paradigm commonly referred to as "vibe coding" \u2014 a term coined by Andrej Karpathy in February 2025 to describe the approach where developers rely almost entirely on AI to produce code without deep engagement with the implementation details.', indent=True)

doc.add_paragraph()

para('However, empirical evidence has revealed significant limitations in single-model code generation. Research conducted in late 2025 and early 2026 has shown that AI-generated code exhibits 1.7x more major issues than human-written code (CodeRabbit, December 2025), 24.7% of AI-generated code contains security flaws, and security vulnerabilities are 2.74x higher in AI co-authored code. Furthermore, a randomized controlled trial in 2025 found that developers actually take 19% longer when using AI tools on projects they know well, suggesting that the perceived productivity gains may be illusory for experienced developers.', indent=True)

doc.add_paragraph()

para('The fundamental problem underlying these findings is the absence of quality assurance in the code generation pipeline. When a single AI model generates code, there is no independent review, no automated testing, and no structured validation of the output. The user receives whatever the model produces and must manually verify its correctness \u2014 a process that defeats the purpose of AI-assisted development for non-expert users and introduces hidden risks for all users.', indent=True)

doc.add_paragraph()

doc.add_heading('1.2 Multi-Agent Systems for Software Engineering', level=2)

para('Multi-agent systems (MAS) offer a promising alternative to single-model approaches. By decomposing the software development process into distinct roles handled by specialized agents, MAS can introduce the same division of labour, review processes, and quality gates that characterize professional software engineering teams. Research systems such as MetaGPT (ICLR 2024), ChatDev (ACL 2024), and AgentCoder (arXiv 2024) have demonstrated that multi-agent architectures consistently outperform single-model approaches on code generation benchmarks.', indent=True)

doc.add_paragraph()

para('MetaGPT introduced the concept of Standardized Operating Procedures (SOPs) for agent communication, requiring agents to communicate through structured documents rather than free-form dialogue. This approach reduced human revision requirements from 2.5 per project to 0.83. AgentCoder demonstrated that separating test generation from code generation \u2014 so that tests are written against specifications rather than implementations \u2014 achieves 96.3% Pass@1 on HumanEval while using 2.4x fewer tokens than MetaGPT. These results establish a clear empirical case for structured, multi-agent approaches to code generation.', indent=True)

doc.add_paragraph()

doc.add_heading('1.3 Co-Lab AI: An Overview', level=2)

para('Co-Lab AI is a platform that applies these research findings to the practical problem of full-stack web application generation. The system employs six specialized AI agents organized in a sequential pipeline with parallel execution where possible:', indent=True)

doc.add_paragraph()

bullet('Understanding Agent: Analyzes the user\'s natural language request, identifies ambiguities, and generates clarifying questions when necessary.')
bullet('Orchestrator Agent: Plans the application architecture, selects the technology stack, defines a formal API contract, and decomposes the work into frontend and backend tasks with individual complexity ratings.')
bullet('Frontend Code Agent: Generates the complete frontend codebase (React, TypeScript, Tailwind CSS) following the API contract specification.')
bullet('Backend Code Agent: Generates the complete backend codebase (Node.js, Express) following the same API contract, running in parallel with the Frontend Agent.')
bullet('Review Agent: Verifies feature completeness, API compatibility between frontend and backend, security posture, and code quality.')
bullet('Test Agent: Generates an independent test suite against the specification (not the implementation) to prevent test bias, producing coverage metrics across endpoints, features, and security dimensions.')

doc.add_paragraph()

para('Additionally, a Quality Scorer function computes a weighted grade (A through F) across five metrics, and a Feedback Fix Agent automatically remediates low-quality output by re-running only the affected agents with targeted fix instructions.', indent=True)

doc.add_paragraph()

doc.add_heading('1.4 Scope of the Project', level=2)

para('This project encompasses the design, implementation, and evaluation of the complete Co-Lab AI platform, including:', indent=True)

doc.add_paragraph()

bullet('The multi-agent AI pipeline with six specialized agents and a feedback loop.')
bullet('A real-time web interface with WebSocket-based streaming, an embedded code editor (Monaco), and in-browser code execution (WebContainers).')
bullet('Support for five AI providers with per-agent model configuration.')
bullet('A quality assurance system with automated grading and selective remediation.')
bullet('Token efficiency optimizations informed by research on trajectory reduction, prompt compression, and agent dropout.')
bullet('MCP (Model Context Protocol) server integration for extensible tool usage.')
bullet('A permission system for governing agent and provider access.')

doc.add_paragraph()

para('The platform is designed as a monorepo using NPM Workspaces and Turborepo, with the frontend (apps/web) built on React 19 and Next.js 16, and the backend (apps/api) built on Express 5 with MongoDB for persistence and WebSockets for real-time communication.', indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  3. MOTIVATION (pages 6-7)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('2. Motivation', level=1)

doc.add_heading('2.1 The Limitations of Current AI Coding Tools', level=2)

para('The motivation for Co-Lab AI stems from the observation that existing AI coding tools fall into three categories, each with significant limitations:', indent=True)

doc.add_paragraph()

para('General-Purpose AI Chatbots (ChatGPT, Claude, Gemini)', bold=True)
para('These tools generate code within a single conversation context. The user provides a prompt, receives a code response, and must manually verify its correctness. There is no formal specification, no independent review, and no testing. When generating full-stack applications, the frontend and backend are produced sequentially or in a single pass, with no guarantee that API calls in the frontend match the routes implemented in the backend. The user receives a monolithic code blob and must debug integration issues themselves.', indent=True)

doc.add_paragraph()

para('AI App Builders (Bolt.new, Lovable, v0, Replit Agent)', bold=True)
para('These platforms provide a more polished interface but fundamentally use a single AI model behind the scenes. They offer no quality grading, no independent test generation, and no automatic feedback loops. The user cannot see what the AI "thought" during generation, cannot review or modify the architectural plan before code generation begins, and cannot choose different models for different tasks. When the output is poor, the only recourse is to re-prompt manually.', indent=True)

doc.add_paragraph()

para('Autonomous Coding Agents (Devin, Claude Code, Codex)', bold=True)
para('These tools are designed for professional developers working on existing codebases. They edit files, run tests, and make pull requests within real repositories. However, they are not designed for generating complete new applications from scratch, and they are not accessible to non-developers who want to create web applications.', indent=True)

doc.add_paragraph()

doc.add_heading('2.2 The Gap We Identified', level=2)

para('None of the existing tools combine the following capabilities in a single platform:', indent=True)

doc.add_paragraph()

numbered('A multi-agent pipeline where different AI models handle different roles simultaneously.')
numbered('A formal API contract that guarantees frontend-backend compatibility before code generation begins.')
numbered('Independent test generation that validates the output against specifications, not implementations.')
numbered('Automated quality grading with measurable metrics and letter grades.')
numbered('A feedback loop that automatically fixes broken output without user intervention.')
numbered('Transparency into the AI\'s reasoning at every stage, with human-in-the-loop checkpoints.')
numbered('In-browser code execution for instant previews without local setup.')

doc.add_paragraph()

para('Co-Lab AI was built to fill this gap \u2014 to create a platform where AI code generation is treated as a team effort with built-in quality assurance, not a single-shot generation with no validation.', indent=True)

doc.add_paragraph()

doc.add_heading('2.3 Academic Motivation', level=2)

para('From an academic perspective, the project is motivated by the growing body of research on multi-agent systems for software engineering. A survey by Fan et al. (ACM TOSEM 2024) identified key gaps in the field: no existing system handles full-stack web application generation end-to-end, no system combines planning, coding, reviewing, and testing agents in a single pipeline, and no system provides human-in-the-loop verification at multiple stages. Co-Lab AI addresses all three of these identified gaps.', indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  4. LITERATURE REVIEW (pages 8-10)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('3. Literature Review / Comparative Work Evaluation', level=1)

doc.add_heading('3.1 Multi-Agent Architectures for Code Generation', level=2)

para('MetaGPT (Hong et al., ICLR 2024) introduced the concept of meta-programming for multi-agent collaboration, where agents communicate through structured documents following Standardized Operating Procedures. The system assigns roles such as Product Manager, Architect, and Engineer to different agents, each producing formal artifacts (PRDs, architecture designs, code). MetaGPT achieved state-of-the-art results on software engineering benchmarks while reducing the need for human revisions by 67%. Co-Lab AI adopts MetaGPT\'s principle of structured communication through its API contract specification, which serves as the single source of truth shared between the Frontend and Backend agents.', indent=True)

doc.add_paragraph()

para('ChatDev (Qian et al., ACL 2024) proposed a chat-chain approach where agents engage in role-playing dialogues to develop software. The system decomposes the development process into sequential phases (design, coding, testing, documenting), with each phase conducted as a conversation between two agents. ChatDev introduced "communicative dehallucination" \u2014 using structured dialogue to reduce fabricated outputs. While Co-Lab AI does not use role-playing dialogue, it adopts ChatDev\'s phase decomposition principle and extends it with parallel execution of independent phases.', indent=True)

doc.add_paragraph()

para('AgentCoder (Huang et al., arXiv 2024) demonstrated the critical importance of separating test generation from code generation. The system uses three agents: a programmer, a test designer, and a test executor. The test designer never sees the programmer\'s output, generating tests solely from the specification. This prevents "test hallucination bias" where tests are inadvertently written to match buggy code. AgentCoder achieved 96.3% Pass@1 on HumanEval with 56.9K token overhead, compared to MetaGPT\'s 85.9% with 138.2K tokens. Co-Lab AI directly implements this principle in its Test Agent, which receives only the API contract and feature list, never the generated code.', indent=True)

doc.add_paragraph()

doc.add_heading('3.2 Orchestration and Dynamic Routing', level=2)

para('The Evolving Orchestration framework (NeurIPS 2025) addressed the rigidity of static agent pipelines by training an orchestrator via reinforcement learning to dynamically select which agents to activate based on evolving task state. The framework achieved 12-25% improvements over static baselines. Co-Lab AI implements a simplified version of this concept through rule-based dynamic routing: the pipeline adapts its behavior based on the detected intent (build/iterate/debug) and the assessed complexity (1-5 scale), skipping unnecessary stages for simple tasks.', indent=True)

doc.add_paragraph()

para('MAGIS (Tao et al., NeurIPS 2024) introduced complexity-aware task decomposition for GitHub issue resolution. Their Manager agent decomposes complex issues into smaller tasks, and the system was shown to reduce the correlation between complexity and failure rate from -25.15 to -1.55. Co-Lab AI adopts this complexity scoring approach, with the Orchestrator rating overall project complexity and individual task complexity on a 1-5 scale.', indent=True)

doc.add_paragraph()

doc.add_heading('3.3 Quality Assurance and Hallucination Reduction', level=2)

para('CodeAgent (Zhang et al., EMNLP 2024) proposed a multi-agent framework for code review with a supervisory QA-Checker agent. The QA-Checker monitors agent conversations for "prompt drifting" and scores quality across multiple dimensions. This approach improved vulnerability detection from 51.42% (GPT-4 alone) to 92.96%. Co-Lab AI\'s quality scoring system is directly inspired by this pattern, using a supervisory function that evaluates output across five weighted metrics.', indent=True)

doc.add_paragraph()

para('Chain-of-Verification (CoVe, Meta Research 2023) demonstrated that having an AI verify its own answers through independent verification questions significantly reduces factual errors. Co-Lab AI embeds CoVe-inspired self-verification steps directly into the code agent prompts, requiring agents to verify API contract compliance before outputting code.', indent=True)

doc.add_paragraph()

doc.add_heading('3.4 Token Efficiency', level=2)

para('AgentDiet (FSE 2026) discovered that 40-60% of agent conversation history is waste \u2014 content that is useless, redundant, or expired. By automatically removing this waste, AgentDiet saved 39.9-59.7% of input tokens. Co-Lab AI implements trajectory reduction by providing downstream agents with compressed context summaries rather than full upstream outputs.', indent=True)

doc.add_paragraph()

para('AgentDropout (ACL 2025) showed that not all agents contribute equally in every round. By applying "dropout" to the agent communication graph, removing agents with minimal contribution, the system saved 21.6% of prompt tokens while improving performance. Co-Lab AI implements this in its feedback loop, where only the agent responsible for the identified issues is re-executed.', indent=True)

doc.add_paragraph()

doc.add_heading('3.5 Comparative Summary', level=2)

add_table(
    ['System', 'Agents', 'API Contract', 'Independent Tests', 'Quality Grading', 'Feedback Loop'],
    [
        ['MetaGPT', '5', 'No', 'No', 'No', 'Limited'],
        ['ChatDev', '4+', 'No', 'No', 'No', 'No'],
        ['AgentCoder', '3', 'No', 'Yes', 'No', 'Yes'],
        ['CodeAgent', '3', 'No', 'No', 'Yes (review)', 'No'],
        ['Co-Lab AI', '6+2', 'Yes', 'Yes', 'Yes (A-F)', 'Yes (selective)'],
    ],
    caption='Comparative evaluation of multi-agent code generation systems',
    table_num=1
)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  5. GAP ANALYSIS (pages 11-13)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('4. Gap Analysis', level=1)

doc.add_heading('4.1 Gaps in Existing Multi-Agent Systems', level=2)

para('Through our comprehensive review of the literature and existing tools, we identified the following critical gaps that Co-Lab AI addresses:', indent=True)

doc.add_paragraph()

para('Gap 1: No Full-Stack Web Application Generation', bold=True)
para('Existing multi-agent systems (MetaGPT, ChatDev, AgentCoder) focus on generating individual programs, algorithms, or single-file solutions. None of them handle the generation of complete full-stack web applications comprising multiple frontend components, backend routes, database schemas, and the integration layer between them. The complexity of full-stack development \u2014 where the frontend and backend must agree on API endpoints, request/response formats, authentication mechanisms, and data models \u2014 requires a fundamentally different approach than generating isolated code files.', indent=True)

doc.add_paragraph()

para('Gap 2: No Formal Cross-Layer Contract', bold=True)
para('In single-model code generation, the AI generates frontend and backend code sequentially or in a single pass. There is no explicit specification ensuring that the frontend\'s API calls match the backend\'s implemented routes. This leads to a common failure mode where the frontend calls GET /api/todos but the backend implements GET /api/tasks, or where the frontend sends { title: "..." } but the backend expects { name: "..." }. These mismatches are syntactically valid on each side but produce a non-functional application. Co-Lab AI addresses this with a formal API contract that both code agents must implement exactly, verified by the Review Agent.', indent=True)

doc.add_paragraph()

para('Gap 3: No Independent Quality Measurement', bold=True)
para('No existing system provides a quantified quality assessment of its output. Users of ChatGPT, Bolt.new, and Lovable receive generated code with no indication of its quality, completeness, or security posture. They must manually review everything. Co-Lab AI\'s quality scoring system provides a letter grade (A through F) computed from five weighted metrics (completeness, security, API compatibility, code quality, test coverage), giving users an immediate, objective assessment of the generated output.', indent=True)

doc.add_paragraph()

para('Gap 4: No Selective Automatic Remediation', bold=True)
para('Existing feedback loops (AgentCoder, MetaGPT) re-run entire stages when issues are detected. This is wasteful when the problem is localized to one side (e.g., a missing backend endpoint). Co-Lab AI\'s feedback loop classifies issues as frontend or backend, then selectively re-runs only the affected agent with targeted fix instructions \u2014 saving approximately 50% of the tokens that a full re-run would consume.', indent=True)

doc.add_paragraph()

para('Gap 5: No Human-in-the-Loop at Planning Stage', bold=True)
para('Most multi-agent systems execute their pipeline autonomously from start to finish. The user provides a prompt and receives the final output with no opportunity to review or modify the plan before expensive code generation begins. This leads to wasted computation when the AI misinterprets the requirements. Co-Lab AI introduces two human checkpoints: after requirement understanding (the user confirms the AI\'s interpretation) and after architectural planning (the user reviews and approves the plan, tech stack, and API contract before code generation begins).', indent=True)

doc.add_paragraph()

para('Gap 6: No In-Browser Execution', bold=True)
para('Existing code generation tools require the user to download the generated code, install dependencies, and run it locally to verify that it works. This creates a significant barrier for non-technical users and adds friction for developers. Co-Lab AI embeds WebContainers \u2014 a browser-based Node.js runtime \u2014 that allows users to preview the generated application directly in their browser with zero local setup.', indent=True)

doc.add_paragraph()

doc.add_heading('4.2 Gap Summary Matrix', level=2)

add_table(
    ['Gap', 'ChatGPT/Claude', 'Bolt/Lovable', 'MetaGPT', 'AgentCoder', 'Co-Lab AI'],
    [
        ['Full-stack generation', 'Partial', 'Yes', 'No', 'No', 'Yes'],
        ['API contract', 'No', 'No', 'No', 'No', 'Yes'],
        ['Quality grading', 'No', 'No', 'No', 'No', 'Yes (A-F)'],
        ['Independent tests', 'No', 'No', 'No', 'Yes', 'Yes'],
        ['Selective feedback', 'No', 'No', 'Limited', 'Full re-run', 'Selective'],
        ['Human-in-the-loop', 'N/A', 'No', 'No', 'No', 'Yes (2 checkpoints)'],
        ['In-browser preview', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Multi-provider AI', 'No', 'No', 'No', 'No', 'Yes (5 providers)'],
    ],
    caption='Gap analysis matrix across existing tools and Co-Lab AI',
    table_num=2
)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  6. PROBLEM STATEMENT (page 14)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('5. Problem Statement', level=1)

para('Current AI-powered code generation tools rely on single-model, single-pass architectures that produce full-stack web applications without formal API contracts, independent quality verification, or automated remediation of defects. This results in:', indent=True)

doc.add_paragraph()

numbered('Frontend-backend incompatibility: API calls in the frontend do not match routes implemented in the backend, producing syntactically valid but functionally broken applications.')
numbered('Absence of quality measurement: Users receive generated code with no objective assessment of its completeness, security, or correctness, requiring manual review of every output.')
numbered('No independent validation: When tests are generated, they are written by the same model that produced the code, leading to test bias where tests inadvertently validate buggy implementations.')
numbered('Wasteful error recovery: Feedback loops re-run entire pipelines even when issues are localized to a single component, consuming unnecessary compute and API tokens.')
numbered('No human oversight of planning: The AI\'s architectural decisions are not visible to the user until after expensive code generation has completed, leading to wasted computation on misinterpreted requirements.')

doc.add_paragraph()

para('The problem this project addresses is: How can we design a multi-agent AI system that generates complete, compatible, and quality-assured full-stack web applications with transparent planning, independent verification, and efficient automatic remediation?', bold=True, indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  7. OBJECTIVES (pages 15-16)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('6. Objectives', level=1)

para('The primary objective of this project is to design, implement, and evaluate a multi-agent AI platform for full-stack web application generation that addresses the identified gaps in existing systems. The specific objectives are:', indent=True)

doc.add_paragraph()

doc.add_heading('6.1 Primary Objectives', level=2)

numbered('Design and implement a pipeline of six specialized AI agents (Understanding, Orchestrator, Frontend, Backend, Review, Test) that collaborate to generate complete full-stack web applications from natural language descriptions.')
numbered('Introduce a formal API contract specification generated by the Orchestrator Agent that both the Frontend and Backend agents must implement, ensuring cross-layer compatibility.')
numbered('Implement independent test generation following the AgentCoder principle, where the Test Agent generates tests against the specification without access to the generated code.')
numbered('Develop a quality scoring system that evaluates generated output across five weighted metrics (completeness, security, API compatibility, code quality, test coverage) and assigns a letter grade (A through F).')
numbered('Build an automated feedback loop that classifies issues by component, selectively re-runs only the affected agent, and re-evaluates quality after fixes.')

doc.add_paragraph()

doc.add_heading('6.2 Secondary Objectives', level=2)

numbered('Provide human-in-the-loop checkpoints at the understanding and planning stages, allowing users to review and approve the AI\'s interpretation and architectural plan before code generation.')
numbered('Support five AI providers (OpenAI, Anthropic, Google Gemini, OpenRouter, GLM) with per-agent model configuration, enabling users to assign different models to different roles.')
numbered('Embed an in-browser execution environment using WebContainers that allows users to preview generated applications without local setup.')
numbered('Implement token efficiency optimizations including trajectory reduction, prompt compression, and selective agent re-execution, targeting 35-45% reduction in token consumption.')
numbered('Deliver a real-time user experience with WebSocket-based streaming that shows agent activity, file creation, and code generation token by token.')
numbered('Integrate MCP (Model Context Protocol) server support for extensible tool usage by the AI agents.')

doc.add_paragraph()

doc.add_heading('6.3 Evaluation Objectives', level=2)

numbered('Validate the pipeline through end-to-end benchmark scenarios covering simple applications, complex multi-feature applications, and iterative feature additions.')
numbered('Measure quality scores and demonstrate that the grading system produces realistic assessments that correlate with actual code quality.')
numbered('Quantify token efficiency improvements relative to the unoptimized baseline.')

section_break()

# ═══════════════════════════════════════════════════════════════════
#  8. METHODOLOGY (pages 17-24)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('7. Methodology', level=1)

doc.add_heading('7.1 System Architecture', level=2)

para('Co-Lab AI follows a modular, layered architecture organized as an NPM Workspaces monorepo with two primary workspaces:', indent=True)

doc.add_paragraph()

bullet('apps/web: The frontend client built on React 19, Next.js 16 (App Router), TypeScript, and Tailwind CSS 4.')
bullet('apps/api: The backend server built on Node.js, Express 5, TypeScript, with MongoDB for persistence and WebSockets for real-time communication.')

doc.add_paragraph()

para('The two workspaces communicate through two channels: HTTP REST for standard CRUD operations (authentication, project management, settings) and WebSocket for the real-time agent pipeline. The WebSocket channel streams agent output token-by-token to the frontend, providing immediate visual feedback during generation.', indent=True)

doc.add_paragraph()

doc.add_heading('7.2 The Multi-Agent Pipeline', level=2)

para('The core methodology is a sequential pipeline of specialized agents with parallel execution where phases are independent. The pipeline proceeds through the following phases:', indent=True)

doc.add_paragraph()

para('Phase 1: Understanding', bold=True)
para('The Understanding Agent receives the user\'s natural language prompt and produces a structured summary of the request. It identifies ambiguous decisions (authentication type, database choice, framework preferences, real-time requirements) and generates 0-5 clarifying questions. If questions are generated, the system pauses for user input before proceeding. This agent uses a fast model (Gemini Flash via OpenRouter) to minimize latency.', indent=True)

doc.add_paragraph()

para('Phase 2: Planning (Orchestration)', bold=True)
para('The Orchestrator Agent receives the confirmed understanding and produces a comprehensive TaskFile containing: intent classification (build/iterate/debug), project metadata, technology stack selection, a formal API contract with every endpoint specification, data models, authentication scheme, frontend and backend task breakdowns with individual complexity ratings, and an overall complexity score (1-5). The API contract is the critical artifact \u2014 it defines every endpoint\'s path, HTTP method, request body fields, and response shape. Both code agents must implement this contract exactly.', indent=True)

doc.add_paragraph()

para('A plan validation checkpoint (inspired by the Checkpoint Architecture paper) runs after orchestration. This is a pure code-level check (zero LLM calls) that verifies: every API endpoint has a corresponding task, authentication requirements are reflected in the task list, and there is balance between frontend and backend tasks. If inconsistencies are found, the Orchestrator is re-prompted with specific fix instructions.', indent=True)

doc.add_paragraph()

para('Phase 3: Code Generation (Parallel)', bold=True)
para('The Frontend Code Agent and Backend Code Agent execute in parallel, each receiving the TaskFile with the API contract. Both agents stream their output over the WebSocket connection, allowing the user to see code appearing in real-time. Each agent includes a self-verification step (inspired by Chain-of-Verification) where it checks its own API calls/routes against the contract before finalizing output. The output format is a JSON object mapping file paths to complete file contents.', indent=True)

doc.add_paragraph()

para('For iterate intent (modifying an existing application), the agents receive the previous code snapshot as context. Trajectory reduction (inspired by AgentDiet) ensures that only the most relevant files are included, capped at 6 files filtered by entry points and files mentioned in current tasks.', indent=True)

doc.add_paragraph()

para('Phase 4: Review', bold=True)
para('The Review Agent receives the generated code along with the TaskFile and produces a structured evaluation covering: completion status (frontend and backend generation success, missing features), API compatibility assessment (whether frontend API calls match backend routes), security audit (critical issues vs. best-practice suggestions), a setup guide, and an initial quality score with metric breakdown. The Review Agent receives compressed context (file summaries and extracted imports/routes rather than full source code) to reduce token consumption.', indent=True)

doc.add_paragraph()

para('Phase 5: Testing', bold=True)
para('The Test Agent generates test metadata (not full test source code, following the CodeAgents compact representation principle) against the API contract and feature list. It produces four categories of tests: basic functionality, edge cases, integration tests, and security tests. It also reports coverage metrics for endpoint coverage, feature coverage, and security coverage, along with contract validation identifying which endpoints are covered and which are missing.', indent=True)

doc.add_paragraph()

para('Phase 6: Quality Scoring', bold=True)
para('The computeQualityScore function combines the Review Agent\'s assessment with count-based heuristics to produce a final score. Five metrics are computed with the following weights:', indent=True)

doc.add_paragraph()

add_table(
    ['Metric', 'Weight', 'Calculation Method'],
    [
        ['Completeness', '25%', 'Based on frontend/backend generation success and missing feature count'],
        ['Security', '20%', 'Penalty per critical issue (x10), capped, floor at 30'],
        ['API Compatibility', '25%', 'Based on compatible flag + mismatch count'],
        ['Code Quality', '15%', 'Base 80, adjusted by issue count and missing features'],
        ['Test Coverage', '15%', 'Average of endpoint, feature, and security coverage percentages'],
    ],
    caption='Quality scoring metrics and weights',
    table_num=3
)

para('The overall score is the average of the formula-based score and the Review Agent\'s self-reported score, blending structural rigour with contextual AI judgment. Letter grades are assigned as: A (90+), B (80-89), C (70-79), D (60-69), F (<60).', indent=True)

doc.add_paragraph()

para('Phase 7: Feedback Loop (Conditional)', bold=True)
para('The feedback loop triggers only when the overall score is below 65 AND there is evidence of real breakage (API incompatible, both sides failed to generate, or 3+ critical security issues). When triggered, issues are classified as frontend or backend using regex pattern matching, and only the affected agent is re-executed with targeted fix instructions. The fixed code is merged with the original (not regenerated from scratch), and quality is re-scored. Maximum one feedback iteration is performed to prevent infinite loops.', indent=True)

doc.add_paragraph()

doc.add_heading('7.3 Real-Time Communication Architecture', level=2)

para('The WebSocket server implements a stateful pipeline with session management, event buffering, and reconnection support:', indent=True)

doc.add_paragraph()

bullet('Session Management: Each WebSocket connection receives a unique session ID. The connection context tracks pipeline state, abort controllers, and an event buffer.')
bullet('Event Buffer: A circular buffer of the last 100 events enables reconnection replay. When a client reconnects after a network interruption, it sends its last received sequence number, and the server replays missed events.')
bullet('Pipeline Persistence: On disconnection, the current pipeline state is saved to a PipelineRun document in MongoDB, enabling the pipeline to resume on reconnection rather than restarting.')
bullet('Heartbeat: 30-second ping/pong intervals detect dead connections.')
bullet('SSE Fallback: For environments where WebSocket connections are blocked, the system falls back to Server-Sent Events with REST-based command submission.')
bullet('Pipeline Queue: If a user submits a new request while one is running, it is queued (per-project FIFO) and processed after the current pipeline completes.')

doc.add_paragraph()

doc.add_heading('7.4 Security Architecture', level=2)

para('The platform implements multiple security layers:', indent=True)

doc.add_paragraph()

bullet('Authentication: JWT tokens stored in httpOnly, secure, SameSite cookies prevent XSS and CSRF attacks. Google OAuth is supported as an alternative login method.')
bullet('API Key Encryption: User API keys are encrypted at rest using AES-256-GCM with random IVs, stored in the format "enc:iv:tag:ciphertext".')
bullet('Rate Limiting: Authentication endpoints are limited to 15 requests per 15 minutes; general API endpoints to 100 requests per minute.')
bullet('Input Validation: All request bodies are validated against Zod schemas before processing.')
bullet('MCP Security: The MCP client manager maintains a command allowlist (npx, node, python, python3, uvx, bunx) for stdio transport, preventing arbitrary command execution.')
bullet('Permission System: A rule-based permission system governs agent access, provider access, token budgets, and feedback loop activation at both user and project levels.')

doc.add_paragraph()

doc.add_heading('7.5 Token Efficiency Optimizations', level=2)

para('Multiple optimizations reduce token consumption without degrading output quality:', indent=True)

doc.add_paragraph()

add_table(
    ['Optimization', 'Research Basis', 'Measured Savings'],
    [
        ['Trajectory reduction for Review/Test agents', 'AgentDiet (FSE 2026)', '58% reduction in Review Agent tokens'],
        ['Test Agent metadata-only mode', 'CodeAgents (arXiv 2025)', '76% reduction in Test Agent tokens'],
        ['System prompt distillation', 'Prompt Compression Survey (NAACL 2025)', '40% reduction in system prompt tokens'],
        ['API contract deduplication', 'AgentDiet', '75% reduction in contract tokens'],
        ['Selective agent re-execution', 'AgentDropout (ACL 2025)', '50% savings per feedback loop'],
        ['Plan validation checkpoint', 'Checkpoint Architecture (arXiv 2026)', 'Prevents 16K-35K token feedback loops'],
        ['Context budget management', '4-tier overflow recovery', 'Prevents prompt truncation'],
    ],
    caption='Token efficiency optimizations with measured savings',
    table_num=4
)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  9. TOOLS / PLATFORM USED (pages 25-26)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('8. Tools / Platform Used', level=1)

doc.add_heading('8.1 Frontend Technologies', level=2)

add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['React', '19.1', 'UI framework with hooks and functional components'],
        ['Next.js', '16.0', 'React framework with App Router, SSR, and API rewrites'],
        ['TypeScript', '5.8', 'Static type safety across the entire frontend codebase'],
        ['Tailwind CSS', '4.1', 'Utility-first CSS framework with dark theme and glass-morphism'],
        ['GSAP', '3.14', 'Animation library for complex effects (orbital, staggered entrance)'],
        ['Monaco Editor', '4.7', 'VS Code-quality code editor embedded in the browser IDE'],
        ['WebContainer API', '1.6', 'Browser-based Node.js runtime for live application preview'],
        ['Radix UI', 'Various', '40+ headless accessible UI component primitives'],
        ['Axios', '1.12', 'HTTP client with auth interceptors'],
        ['Prism React Renderer', '2.4', 'Syntax highlighting for code blocks (Night Owl theme)'],
        ['Recharts', '3.2', 'Data visualization for quality metrics and benchmarks'],
        ['Lucide React', '0.544', 'Consistent SVG icon library'],
        ['React Hook Form', '7.63', 'Form state management with validation'],
        ['Three.js', '0.183', '3D graphics for visual effects'],
    ],
    caption='Frontend technology stack',
    table_num=5
)

doc.add_heading('8.2 Backend Technologies', level=2)

add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Node.js', '18+', 'Server-side JavaScript runtime'],
        ['Express', '5.1', 'HTTP framework with native Promise support in route handlers'],
        ['TypeScript', '5.8', 'Type safety for the backend codebase'],
        ['Mongoose', '8.18', 'MongoDB ODM with schema validation and population'],
        ['ws', '8.18', 'WebSocket server for real-time agent pipeline streaming'],
        ['OpenAI SDK', '5.22', 'Client for OpenAI, OpenRouter, and GLM (OpenAI-compatible APIs)'],
        ['Anthropic SDK', '0.78', 'Client for Claude models'],
        ['Google GenAI', '1.42', 'Client for Gemini models'],
        ['Zod', '3.25', 'Runtime schema validation for API inputs'],
        ['jsonwebtoken', '9.0', 'JWT-based stateless authentication'],
        ['bcryptjs', '3.0', 'Password hashing with configurable salt rounds'],
        ['MCP SDK', '1.12', 'Model Context Protocol client for tool integration'],
        ['express-rate-limit', '8.3', 'Rate limiting for API endpoints'],
    ],
    caption='Backend technology stack',
    table_num=6
)

doc.add_heading('8.3 Infrastructure', level=2)

add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Database', 'MongoDB Atlas', 'Primary data persistence (13 collections)'],
        ['Dynamic DB Provisioning', 'Turso (SQLite)', 'Per-project database provisioning via REST API'],
        ['Build Orchestration', 'Turborepo 2.4', 'Monorepo build coordination'],
        ['Package Management', 'NPM Workspaces', 'Workspace-aware dependency management'],
        ['Version Control', 'Git', 'Source code version control'],
    ],
    caption='Infrastructure components',
    table_num=7
)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  10. ENVIRONMENTAL SETUP (pages 27-28)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('9. Environmental Setup', level=1)

doc.add_heading('9.1 Prerequisites', level=2)

numbered('Node.js version 18 or higher installed on the development machine.')
numbered('NPM (Node Package Manager) included with Node.js.')
numbered('A MongoDB Atlas account with a cluster URI, or a local MongoDB instance.')
numbered('API keys for at least one AI provider (OpenAI, Anthropic, Google Gemini, OpenRouter, or GLM).')
numbered('(Optional) A Turso account for dynamic SQLite database provisioning.')
numbered('(Optional) A Google Cloud project with OAuth credentials for Google Sign-In.')

doc.add_paragraph()

doc.add_heading('9.2 Installation Steps', level=2)

para('Step 1: Clone the repository', bold=True)
p = doc.add_paragraph()
run = p.add_run('git clone <repository-url>\ncd Co-Lab-AI')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()

para('Step 2: Install dependencies', bold=True)
p = doc.add_paragraph()
run = p.add_run('npm install')
run.font.name = 'Courier New'
run.font.size = Pt(10)

para('This command installs dependencies for both the frontend (apps/web) and backend (apps/api) workspaces simultaneously, as configured in the root package.json workspaces field.', indent=True)

doc.add_paragraph()

para('Step 3: Configure the backend environment', bold=True)
para('Create a .env file in apps/api/ with the following required variables:', indent=True)

doc.add_paragraph()

add_table(
    ['Variable', 'Required', 'Description'],
    [
        ['DATABASE_URL', 'Yes', 'MongoDB connection URI'],
        ['JWT_SECRET', 'Yes', 'Secret key for JWT token signing'],
        ['ENCRYPTION_KEY', 'Production', '64-character hex string for AES-256-GCM API key encryption'],
        ['OPENROUTER_API_KEY', 'Recommended', 'OpenRouter API key (access to multiple models)'],
        ['OPENAI_API_KEY', 'Optional', 'Direct OpenAI API key'],
        ['ANTHROPIC_API_KEY', 'Optional', 'Anthropic API key for Claude models'],
        ['GEMINI_API_KEY', 'Optional', 'Google Gemini API key'],
        ['GLM_API_KEY', 'Optional', 'ZhipuAI GLM API key'],
        ['GOOGLE_CLIENT_ID', 'Optional', 'Google OAuth client ID'],
        ['CORS_ORIGINS', 'Optional', 'Comma-separated allowed origins (defaults to localhost:3000)'],
        ['PORT', 'Optional', 'Server port (defaults to 5000)'],
    ],
    caption='Backend environment variables',
    table_num=8
)

para('Step 4: Configure the frontend environment', bold=True)
para('Copy apps/web/.env.example to apps/web/.env.local and set NEXT_PUBLIC_API_URL to the backend URL (defaults to /api/v1 with Next.js rewrites handling the proxy).', indent=True)

doc.add_paragraph()

para('Step 5: Start the development servers', bold=True)
p = doc.add_paragraph()
run = p.add_run('npm run dev')
run.font.name = 'Courier New'
run.font.size = Pt(10)

para('Turborepo starts both the frontend (port 3000) and backend (port 5000) development servers concurrently. The frontend dev server (dev-server.mjs) automatically proxies WebSocket connections to the backend.', indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  11. IMPLEMENTATION (pages 29-32)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('10. Implementation', level=1)

doc.add_heading('10.1 Backend Implementation', level=2)

para('The backend is organized into a modular architecture with clear separation of concerns:', indent=True)

doc.add_paragraph()

para('Agents (agents/): Each of the six pipeline agents is implemented as an individual module with a single exported async function. The Understanding and Orchestrator agents use non-streaming AI calls (callAIGenerate), while the code agents (Frontend, Backend), Review, and Test agents use streaming calls (callAIGenerateStream) that yield chunks via async generators. The helpers module provides shared utilities for trajectory reduction (buildCompactContract, compressSnapshotForAgent), plan validation (validateOrchestratorPlan), and memory decay filtering (filterStaleItems).', indent=True)

doc.add_paragraph()

para('Services (services/): Business logic is separated into dedicated service modules. The AI abstraction layer (ai-clients.ts, ai-generate.ts) provides a unified interface across five providers. The JSON parser (json-parser.ts) implements five strategies for extracting JSON from LLM output with schema validation and automatic retry. The crypto service (crypto.ts) handles AES-256-GCM encryption of API keys at rest. The context budget manager (context-budget.ts) implements four-tier overflow recovery to prevent prompt truncation.', indent=True)

doc.add_paragraph()

para('WebSocket Layer (websocket/): The WebSocket server maintains per-connection state including pipeline phase, event buffers, and abort controllers. A circular event buffer of 100 events enables client reconnection with event replay. Pipeline state is persisted to MongoDB (PipelineRun) on disconnection, allowing sessions to resume. Separate handler modules process each client message type (message, understanding_response, qa_complete, proceed, resume, permission_response).', indent=True)

doc.add_paragraph()

para('Data Models (models/): Thirteen Mongoose schemas define the data layer. The Message schema captures the complete pipeline history for each user request, including responses from every agent, quality scores, and feedback iteration counts. The ProjectSnapshot schema stores the latest generated code per project with a projectMemory field that tracks preferred patterns, rejected approaches, and quality feedback with staleness decay.', indent=True)

doc.add_paragraph()

doc.add_heading('10.2 Frontend Implementation', level=2)

para('The frontend uses Next.js 16 App Router for file-based routing. Protected routes are wrapped with a ProtectedPage component that checks authentication state and redirects unauthenticated users to /login with a return URL.', indent=True)

doc.add_paragraph()

para('The main chat interface (screens/App.tsx, approximately 1100 lines) integrates three core hooks: useWebSocket for real-time pipeline communication, useActivityTracker for converting WebSocket events into structured activity items, and useWebContainer for in-browser code execution. The pipeline status is visualized through multiple components: PipelineInsightBar shows phase progress, AgentActivityPanel shows detailed activity with a mission-control timeline, and ThinkingIndicator provides orbital animations during agent processing.', indent=True)

doc.add_paragraph()

para('The IDE modal (components/IDEModal.tsx) provides a full development environment with Monaco Editor for code viewing, a file tree navigator, resizable panels, a terminal showing build output, and a live preview iframe connected to the WebContainer runtime. The WebContainer boots a complete Node.js environment in the browser, installs dependencies, and starts development servers for both frontend and backend code.', indent=True)

doc.add_paragraph()

para('Real-time streaming is handled by the useWebSocket hook, which manages WebSocket connections with exponential backoff reconnection (1 second base, 30 second maximum, with random jitter). After three consecutive WebSocket failures, the hook automatically falls back to Server-Sent Events via the useSSEFallback hook, ensuring connectivity in restrictive network environments.', indent=True)

doc.add_paragraph()

doc.add_heading('10.3 AI Provider Integration', level=2)

para('The AI abstraction layer supports five providers through a unified interface. For OpenAI-compatible APIs (OpenAI, OpenRouter, GLM), the OpenAI SDK is reused with different base URLs. Anthropic and Google Gemini use their respective native SDKs due to API format differences. The user settings system allows per-agent model configuration, with defaults of Gemini Flash for fast agents (Understanding, Review, Test) and Claude Sonnet for quality-critical agents (Frontend, Backend).', indent=True)

doc.add_paragraph()

doc.add_heading('10.4 MCP Server Integration', level=2)

para('Co-Lab AI supports the Model Context Protocol (MCP) for extensible tool usage. Users can register MCP servers using three transport types: stdio (child process), HTTP-SSE, and streamable-HTTP. The MCPClientManager maintains connections with a 5-minute idle timeout and caches discovered tools. Tool definitions are injected into agent prompts when relevant, allowing agents to leverage external capabilities during code generation.', indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  12. RESULTS AND DISCUSSION (pages 33-35)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('11. Results and Discussion', level=1)

doc.add_heading('11.1 Evaluation Methodology', level=2)

para('The system was evaluated using three end-to-end benchmark scenarios of increasing complexity, executed with production AI models (Claude Sonnet 4.6 for code generation, Gemini 2.5 Flash for planning/review/testing) via OpenRouter:', indent=True)

doc.add_paragraph()

bullet('Scenario 1 (Simple): "Build a simple todo list app" \u2014 basic CRUD, no authentication, in-memory storage.')
bullet('Scenario 2 (Complex): "Build a full e-commerce platform with authentication, products, cart, orders, admin panel, and payments" \u2014 multiple resources, relationships, and security requirements.')
bullet('Scenario 3 (Iterate): "Add a product review and rating system" to the existing e-commerce application \u2014 tests snapshot-based iteration and intent detection.')

doc.add_paragraph()

doc.add_heading('11.2 Quality Assessment Results', level=2)

add_table(
    ['Scenario', 'Complexity', 'Quality Grade', 'Overall Score', 'Tests Generated', 'Time'],
    [
        ['Simple Todo App', '2/5', 'A', '90/100', '32 tests', '95s'],
        ['E-commerce Platform', '4/5', 'C', '78/100', '82 tests', '198s'],
        ['Iterate: Add Reviews', '4/5', 'B', '87/100', '100 tests', '241s'],
    ],
    caption='End-to-end evaluation results',
    table_num=9
)

doc.add_heading('11.3 Per-Metric Breakdown', level=2)

add_table(
    ['Metric', 'Todo App', 'E-commerce', 'Iterate'],
    [
        ['Completeness', '90', '35', '73'],
        ['Security', '85', '88', '88'],
        ['API Compatibility', '98', '95', '95'],
        ['Code Quality', '89', '81', '86'],
        ['Test Coverage', '92', '97', '90'],
    ],
    caption='Per-metric quality scores across benchmark scenarios',
    table_num=10
)

para('The simple todo app achieved grade A (90/100), demonstrating that the system produces high-quality output for well-scoped applications. The e-commerce platform received grade C (78/100), primarily due to frontend generation hitting token limits (completeness scored 35), while the backend was fully implemented. The iterate scenario achieved grade B (87/100), correctly building on the existing snapshot and implementing the review system as requested.', indent=True)

doc.add_paragraph()

doc.add_heading('11.4 Token Efficiency Results', level=2)

add_table(
    ['Component', 'Before Optimization', 'After Optimization', 'Reduction'],
    [
        ['Review Agent output', '4,274 tokens', '1,785 tokens', '58%'],
        ['Test Agent output', '11,431 tokens', '2,763 tokens', '76%'],
        ['System prompts (input)', '~3,290 tokens', '~1,970 tokens', '40%'],
        ['API contract duplication', '~3,200 tokens', '~800 tokens', '75%'],
    ],
    caption='Token efficiency improvements',
    table_num=11
)

para('Overall token consumption was reduced by an estimated 35-45% per pipeline run. The most significant saving came from the Test Agent metadata-only mode (76% reduction), followed by trajectory reduction for the Review Agent (58% reduction).', indent=True)

doc.add_paragraph()

doc.add_heading('11.5 Speed Improvements', level=2)

add_table(
    ['Scenario', 'Before', 'After', 'Speedup'],
    [
        ['Simple Todo', '~670s', '~95s', '7x'],
        ['E-commerce', '~533s', '~198s', '2.7x'],
        ['Iterate', '~255s', '~241s', '~1x (already fast)'],
    ],
    caption='Pipeline execution time improvements',
    table_num=12
)

doc.add_heading('11.6 Discussion', level=2)

para('The results demonstrate that the multi-agent approach with formal API contracts produces measurably better output than single-model generation. The API compatibility scores (95-98 across all scenarios) validate the effectiveness of the contract specification approach \u2014 frontend and backend code is consistently compatible. The independent test generation produced comprehensive coverage (32-100 tests per scenario) that identified issues the Review Agent missed.', indent=True)

doc.add_paragraph()

para('The primary limitation revealed by the evaluation is token-limited generation for complex applications. The e-commerce scenario\'s low completeness score (35) was caused by the frontend agent exhausting its token budget before completing all components. This suggests that for complex applications, a chunked generation approach or adaptive token allocation would be beneficial.', indent=True)

doc.add_paragraph()

para('The quality scoring system\'s calibration proved effective: working applications received high grades (A-B), while applications with genuine issues received proportionally lower grades (C). The feedback loop correctly triggered only on real breakage, avoiding unnecessary re-generation for minor issues.', indent=True)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  13. CONCLUSION & FUTURE WORK (pages 36-37)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('12. Conclusion & Future Work', level=1)

doc.add_heading('12.1 Conclusion', level=2)

para('This project demonstrates that multi-agent AI systems with structured communication, independent verification, and automated quality assurance can produce significantly better full-stack web applications than single-model approaches. The key contributions of Co-Lab AI are:', indent=True)

doc.add_paragraph()

numbered('A practical implementation of a six-agent pipeline that handles the complete software development lifecycle from requirements to tested, graded code.')
numbered('A formal API contract specification that eliminates the most common failure mode in AI-generated full-stack applications: frontend-backend incompatibility.')
numbered('An independent test generation approach (following the AgentCoder principle) that produces comprehensive test suites without implementation bias.')
numbered('A quantified quality assessment system that gives users immediate, objective feedback on generated code quality.')
numbered('Token efficiency optimizations of 35-45% informed by recent research on trajectory reduction, prompt compression, and agent dropout.')
numbered('A production-grade platform with real-time streaming, in-browser code execution, multi-provider AI support, and session resumability.')

doc.add_paragraph()

para('The evaluation across three benchmark scenarios confirmed that the system produces grade-A output for well-scoped applications and gracefully degrades with measurable quality feedback for complex applications. The platform successfully bridges the gap between research multi-agent systems and practical developer tools.', indent=True)

doc.add_paragraph()

doc.add_heading('12.2 Future Work', level=2)

para('Several directions for future development have been identified:', indent=True)

doc.add_paragraph()

numbered('Chunked Generation: For complex applications that exceed token limits, implement a chunked generation strategy where the code agents produce output in multiple passes, each building on the previous output.')
numbered('Reinforcement Learning Orchestration: Replace the current rule-based pipeline routing with a trained orchestrator (as proposed in the Evolving Orchestration paper) that learns optimal agent selection and sequencing from past generation outcomes.')
numbered('Collaborative Editing: Enable multiple users to interact with the same project simultaneously, with real-time synchronization of pipeline state and generated code.')
numbered('Execution-Based Testing: Extend the Test Agent to actually execute generated tests within the WebContainer environment, providing empirical pass/fail results rather than static analysis.')
numbered('AI-Based Cost Estimation: Predict token consumption before pipeline execution based on project complexity, enabling users to make informed decisions about model selection and budget allocation.')
numbered('Expanded Plugin Ecosystem: Deepen integrations with external services (databases, deployment platforms, CI/CD systems) to generate deployment-ready applications.')

section_break()

# ═══════════════════════════════════════════════════════════════════
#  14. REFERENCES (page 38)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('13. References', level=1)

refs = [
    '[1] Hong, S., et al. (2024). "MetaGPT: Meta Programming for a Multi-Agent Collaborative Framework." ICLR 2024. arXiv:2308.00352.',
    '[2] Qian, C., et al. (2024). "ChatDev: Communicative Agents for Software Development." ACL 2024. arXiv:2307.07924.',
    '[3] Huang, D., et al. (2024). "AgentCoder: Multi-Agent-based Code Generation with Iterative Testing and Optimisation." arXiv:2312.13010.',
    '[4] Chen, Y., et al. (2025). "Evolving Orchestration of Multi-Agent Collaboration." NeurIPS 2025. arXiv:2505.19591.',
    '[5] Qian, C., et al. (2025). "Scaling Large-Language-Model-based Multi-Agent Collaboration." ICLR 2025. arXiv:2406.07155.',
    '[6] Tao, W., et al. (2024). "MAGIS: LLM-Based Multi-Agent Framework for GitHub Issue Resolution." NeurIPS 2024. arXiv:2403.17927.',
    '[7] Zhang, H., et al. (2024). "CodeAgent: Autonomous Communicative Agents for Code Review." EMNLP 2024. arXiv:2402.02172.',
    '[8] Fan, A., et al. (2024). "Large Language Model-based Agents for Software Engineering: A Survey." ACM TOSEM. arXiv:2404.04834.',
    '[9] Zhang, Y., et al. (2026). "AgentDiet: Agentic Trajectory Reduction for Efficient LLM Agent Fine-Tuning." FSE 2026. arXiv:2509.23586.',
    '[10] Wang, X., et al. (2025). "AgentDropout: Dynamic Agent Elimination for Token-Efficient Multi-Agent Communication." ACL 2025. arXiv:2503.18891.',
    '[11] Nguyen, T., et al. (2025). "CodeAgents: Building LLM Agents with Codified Multi-Agent Interaction." arXiv:2507.03254.',
    '[12] Li, Z., et al. (2025). "Prompt Compression: A Survey." NAACL 2025 (Oral).',
    '[13] Dhuliawala, S., et al. (2023). "Chain-of-Verification Reduces Hallucination in Large Language Models." Meta AI Research. arXiv:2309.11495.',
    '[14] Anwar, S., et al. (2025). "Rule-Based Multi-Agent System for Hallucination Mitigation." MDPI Information 16(7):517.',
    '[15] Li, Q., et al. (2026). "SupervisorAgent: Autonomous LLM-based Agent for Efficient and Reliable Task Supervision." ICLR 2026. arXiv:2510.26585.',
    '[16] Wang, Z., et al. (2026). "Checkpoint Architecture for Multi-Agent Hallucination Reduction." arXiv:2603.07728.',
    '[17] Han, Y., et al. (2025). "Reducing LLM Hallucination via CoT, RAG, Self-Consistency, and Self-Verification." arXiv:2505.09031.',
    '[18] Chen, L., et al. (2025). "Citation-Grounded Code Comprehension." arXiv:2512.12117.',
    '[19] OPTIMA, et al. (2025). "Optimizing Multi-Agent Communication for Efficiency." ACL 2025.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

section_break()

# ═══════════════════════════════════════════════════════════════════
#  15. ANNEXURE I (page 39)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Annexure I: Plagiarism Declaration Certificate', level=1)

doc.add_paragraph()
doc.add_paragraph()

para('DECLARATION', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

para('We hereby declare that the project entitled "Co-Lab AI: The Multi-Agent AI Platform That Builds Full-Stack Web Applications" submitted to the Department of Computer Science and Engineering is a record of original work done by us under the guidance of our project supervisor.', indent=True)

doc.add_paragraph()

para('The results embodied in this project have not been submitted to any other University or Institute for the award of any degree or diploma.', indent=True)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

para('Date: _______________')
doc.add_paragraph()
para('Devansh')
para('Signature: _______________________')
doc.add_paragraph()
para('Palak')
para('Signature: _______________________')

section_break()

# ═══════════════════════════════════════════════════════════════════
#  16. ANNEXURE II (page 40)
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Annexure II: Plagiarism Report', level=1)

doc.add_paragraph()
doc.add_paragraph()

para('[Attach the plagiarism report generated by Turnitin / iThenticate / other plagiarism detection tool here.]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

para('Note: The plagiarism report should show the similarity index and highlight any matched content sources. The acceptable similarity index is typically below 15-20%.', italic=True, indent=True)


# ─── Add page numbers ──────────────────────────────────────────────
add_page_number(doc)

# ─── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Docs', 'Co-Lab-AI_Major_Project_Report.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
